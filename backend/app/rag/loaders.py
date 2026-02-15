"""
Document loaders for various file formats.

This module provides loaders for:
- Markdown (.md) files
- HTML (.html, .htm) files
- Plain text (.txt) files
- PDF (.pdf) files
- DOCX (.docx) files
"""

import logging
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    logger.warning("pdfplumber not installed. PDF loading will not be available.")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    logger.warning("python-docx not installed. DOCX loading will not be available.")


def load_markdown(file_path: Path) -> dict:
    """
    Load and parse a markdown file.

    Extracts:
    - Title from first H1 heading or filename
    - Full text content
    - Preserves markdown structure

    Args:
        file_path: Path to the markdown file

    Returns:
        Dictionary with keys: 'text', 'title', 'url', 'doc_id'
    """
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            content = file_path.read_text(encoding="latin-1")
        except Exception as e:
            logger.error(f"Failed to read markdown file {file_path}: {e}")
            raise

    # Extract title from first H1 heading
    title_match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()
    else:
        # Fallback to filename without extension
        title = file_path.stem.replace("_", " ").replace("-", " ").title()

    # Generate doc_id from filename
    doc_id = file_path.stem

    # Generate URL (relative path from data/docs)
    url = f"/docs/{file_path.name}"

    return {
        "text": content,
        "title": title,
        "url": url,
        "doc_id": doc_id,
    }


def load_html(file_path: Path) -> dict:
    """
    Load and parse an HTML file.

    Extracts:
    - Title from <title> tag or first <h1> heading
    - Text content with tags stripped, but keeps headings and links
    - Converts headings to markdown-style format
    - Preserves link text and URLs

    Args:
        file_path: Path to the HTML file

    Returns:
        Dictionary with keys: 'text', 'title', 'url', 'doc_id'
    """
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            content = file_path.read_text(encoding="latin-1")
        except Exception as e:
            logger.error(f"Failed to read HTML file {file_path}: {e}")
            raise

    # Extract title from <title> tag
    title_match = re.search(r"<title[^>]*>(.*?)</title>", content, re.IGNORECASE | re.DOTALL)
    if title_match:
        title = re.sub(r"<[^>]+>", "", title_match.group(1)).strip()
    else:
        # Try to extract from first <h1> tag
        h1_match = re.search(r"<h1[^>]*>(.*?)</h1>", content, re.IGNORECASE | re.DOTALL)
        if h1_match:
            title = re.sub(r"<[^>]+>", "", h1_match.group(1)).strip()
        else:
            # Fallback to filename
            title = file_path.stem.replace("_", " ").replace("-", " ").title()

    # Extract text content, preserving headings and links
    # Remove script and style tags
    content = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.IGNORECASE | re.DOTALL)

    # Convert headings to markdown-style
    content = re.sub(r"<h1[^>]*>(.*?)</h1>", r"# \1", content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r"<h2[^>]*>(.*?)</h2>", r"## \1", content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r"<h3[^>]*>(.*?)</h3>", r"### \1", content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r"<h4[^>]*>(.*?)</h4>", r"#### \1", content, flags=re.IGNORECASE | re.DOTALL)

    # Convert links to markdown-style [text](url)
    def link_replacer(match):
        link_text = re.sub(r"<[^>]+>", "", match.group(1))
        href = match.group(2) if match.group(2) else ""
        return f"[{link_text}]({href})"

    content = re.sub(
        r'<a[^>]*href=["\']?([^"\'>\s]*)["\']?[^>]*>(.*?)</a>',
        link_replacer,
        content,
        flags=re.IGNORECASE | re.DOTALL,
    )

    # Remove all remaining HTML tags
    text = re.sub(r"<[^>]+>", "", content)

    # Clean up whitespace
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)  # Multiple blank lines to double
    text = re.sub(r"[ \t]+", " ", text)  # Multiple spaces to single
    text = text.strip()

    # Generate doc_id from filename
    doc_id = file_path.stem

    # Generate URL (relative path from data/docs)
    url = f"/docs/{file_path.name}"

    return {
        "text": text,
        "title": title,
        "url": url,
        "doc_id": doc_id,
    }


def load_txt(file_path: Path) -> dict:
    """
    Load and parse a plain text file.

    Extracts:
    - Title from first line or filename
    - Full text content

    Args:
        file_path: Path to the text file

    Returns:
        Dictionary with keys: 'text', 'title', 'url', 'doc_id'
    """
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            content = file_path.read_text(encoding="latin-1")
        except Exception as e:
            logger.error(f"Failed to read text file {file_path}: {e}")
            raise

    # Extract title from first non-empty line (if it looks like a title)
    lines = content.strip().split("\n")
    title = None
    if lines:
        first_line = lines[0].strip()
        # If first line is short and doesn't end with punctuation, treat as title
        if len(first_line) < 100 and not first_line.endswith((".", "!", "?")):
            title = first_line
            # Remove title from content
            content = "\n".join(lines[1:]).strip()

    if not title:
        # Fallback to filename
        title = file_path.stem.replace("_", " ").replace("-", " ").title()

    # Generate doc_id from filename
    doc_id = file_path.stem

    # Generate URL (relative path from data/docs)
    url = f"/docs/{file_path.name}"

    return {
        "text": content,
        "title": title,
        "url": url,
        "doc_id": doc_id,
    }


def load_pdf(file_path: Path) -> Optional[dict]:
    """
    Load and parse a PDF file.

    Extracts:
    - Title from document metadata or first page heading
    - Text content from all pages
    - Preserves structure (headings, paragraphs)
    - Includes page metadata in text

    Args:
        file_path: Path to the PDF file

    Returns:
        Dictionary with keys: 'text', 'title', 'url', 'doc_id'
        Returns None if PDF library is not available or error occurs
    """
    if pdfplumber is None:
        logger.error("pdfplumber is not installed. Cannot load PDF files.")
        return None

    try:
        text_parts = []
        title = None

        with pdfplumber.open(file_path) as pdf:
            # Try to get title from metadata
            if pdf.metadata and pdf.metadata.get("Title"):
                title = pdf.metadata["Title"].strip()

            # Extract text from each page
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    # Add page marker for multi-page documents
                    if len(pdf.pages) > 1:
                        text_parts.append(f"\n[Page {page_num}]\n")
                    text_parts.append(page_text)

        # Combine all text
        full_text = "\n".join(text_parts).strip()

        if not full_text:
            logger.warning(f"No text content extracted from PDF: {file_path}")
            return None

        # If no title from metadata, try to extract from first page
        if not title:
            first_page_lines = text_parts[0].split("\n")[:5] if text_parts else []
            for line in first_page_lines:
                line = line.strip()
                # Look for a line that looks like a title (short, no punctuation at end)
                if line and len(line) < 100 and not line.endswith((".", "!", "?", ":")):
                    title = line
                    break

        # Fallback to filename if no title found
        if not title:
            title = file_path.stem.replace("_", " ").replace("-", " ").title()

        # Generate doc_id from filename
        doc_id = file_path.stem

        # Generate URL (relative path from data/docs)
        url = f"/docs/{file_path.name}"

        return {
            "text": full_text,
            "title": title,
            "url": url,
            "doc_id": doc_id,
        }

    except Exception as e:
        logger.error(f"Error loading PDF file {file_path}: {e}")
        return None


def load_docx(file_path: Path) -> Optional[dict]:
    """
    Load and parse a DOCX file.

    Extracts:
    - Title from document properties or first heading
    - Text content while preserving structure
    - Handles headings, paragraphs, lists, and tables
    - Converts formatting to markdown-style where appropriate

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary with keys: 'text', 'title', 'url', 'doc_id'
        Returns None if python-docx library is not available or error occurs
    """
    if DocxDocument is None:
        logger.error("python-docx is not installed. Cannot load DOCX files.")
        return None

    try:
        doc = DocxDocument(file_path)

        # Try to get title from document properties
        title = None
        if doc.core_properties.title:
            title = doc.core_properties.title.strip()

        # Extract text content with structure preservation
        text_parts = []

        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                # Empty paragraph - add blank line
                text_parts.append("")
                continue

            para_text = paragraph.text.strip()

            # Check if this is a heading
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                # Extract heading level from style name (e.g., "Heading 1" -> level 1)
                try:
                    level = int(paragraph.style.name.split()[-1])
                    # Convert to markdown heading
                    text_parts.append(f"{'#' * level} {para_text}")
                except (ValueError, IndexError):
                    # Fallback: treat as regular paragraph
                    text_parts.append(para_text)
            else:
                # Regular paragraph
                text_parts.append(para_text)

            # If we haven't found a title yet and this looks like a title, use it
            if not title and para_text and len(para_text) < 100:
                # Check if it's a heading style or first paragraph
                if (paragraph.style and paragraph.style.name.startswith("Heading")) or len(text_parts) == 1:
                    title = para_text

        # Extract text from tables
        for table in doc.tables:
            table_text_parts = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text_parts.append(row_text)
            if table_text_parts:
                text_parts.append("")  # Blank line before table
                text_parts.extend(table_text_parts)
                text_parts.append("")  # Blank line after table

        # Combine all text
        full_text = "\n".join(text_parts).strip()

        if not full_text:
            logger.warning(f"No text content extracted from DOCX: {file_path}")
            return None

        # Fallback to filename if no title found
        if not title:
            title = file_path.stem.replace("_", " ").replace("-", " ").title()

        # Generate doc_id from filename
        doc_id = file_path.stem

        # Generate URL (relative path from data/docs)
        url = f"/docs/{file_path.name}"

        return {
            "text": full_text,
            "title": title,
            "url": url,
            "doc_id": doc_id,
        }

    except Exception as e:
        logger.error(f"Error loading DOCX file {file_path}: {e}")
        return None


def load_document(file_path: Path) -> Optional[dict]:
    """
    Load a document based on its file extension.

    Supports:
    - .md, .markdown -> markdown loader
    - .html, .htm -> HTML loader
    - .txt -> text loader
    - .pdf -> PDF loader
    - .docx -> DOCX loader

    Args:
        file_path: Path to the document file

    Returns:
        Dictionary with keys: 'text', 'title', 'url', 'doc_id'
        Returns None if file format is not supported
    """
    suffix = file_path.suffix.lower()

    if suffix in [".md", ".markdown"]:
        return load_markdown(file_path)
    elif suffix in [".html", ".htm"]:
        return load_html(file_path)
    elif suffix == ".txt":
        return load_txt(file_path)
    elif suffix == ".pdf":
        return load_pdf(file_path)
    elif suffix == ".docx":
        return load_docx(file_path)
    else:
        logger.warning(f"Unsupported file format: {suffix} for file {file_path}")
        return None
